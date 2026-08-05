from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\ai hub\patent-reading-plugin")
IMAGE_DIR = Path(r"C:\Users\talenlin\Desktop\公众号附图")
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "专利阅研使用指南-微信公众号图文版.docx"


# compact_reference_guide preset, with a named product-accent override.
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
CONTENT_WIDTH_DXA = 9360
TEAL = RGBColor(31, 111, 101)
TEAL_DARK = RGBColor(21, 82, 75)
TEAL_LIGHT = "EAF5F2"
BLUE = RGBColor(46, 116, 181)
BLUE_DARK = RGBColor(31, 77, 120)
INK = RGBColor(38, 47, 56)
MUTED = RGBColor(92, 105, 117)
LIGHT_GRAY = "F2F4F7"
WARN_FILL = "FFF4E5"
WARN_TEXT = RGBColor(134, 78, 0)
RED = RGBColor(156, 36, 36)
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=None, bold=None, italic=None, color=None,
                 ascii_font="Calibri", east_asia_font="Microsoft YaHei"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F6F65")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(r_fonts)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_callout(doc, label, text, fill=TEAL_LIGHT, label_color=TEAL_DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA], indent_dxa=120)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"{label}  ")
    set_run_font(r, size=10.5, bold=True, color=label_color)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_number(doc, text, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def add_figure(doc, filename, caption, alt_text, width=6.35, page_break_before=False):
    path = IMAGE_DIR / filename
    p = doc.add_paragraph()
    if page_break_before:
        p.paragraph_format.page_break_before = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.keep_with_next = True
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=MUTED, italic=True)
    return p


def add_section_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, bold=True, color=TEAL_DARK)
    elif level == 2:
        set_run_font(r, size=13, bold=True, color=TEAL)
    else:
        set_run_font(r, size=12, bold=True, color=BLUE_DARK)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = PAGE_WIDTH
section.page_height = PAGE_HEIGHT
section.top_margin = MARGIN
section.right_margin = MARGIN
section.bottom_margin = MARGIN
section.left_margin = MARGIN
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

quick_path_num = create_decimal_numbering(doc)
workflow_num = create_decimal_numbering(doc)

# Styles resolved from compact_reference_guide.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, TEAL_DARK, 18, 10),
    ("Heading 2", 13, TEAL, 14, 7),
    ("Heading 3", 12, BLUE_DARK, 10, 5),
]:
    style = doc.styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

caption_style = doc.styles.add_style("Figure Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
caption_style.font.name = "Calibri"
caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
caption_style.font.size = Pt(9)
caption_style.font.italic = True
caption_style.font.color.rgb = MUTED

# Running header and footer. First-page header intentionally blank for editorial-cover treatment.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("专利阅研｜本地专利阅读与辅助审查")
set_run_font(hr, size=9, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("专利阅研使用指南  ·  ")
set_run_font(fr, size=9, color=MUTED)
add_page_field(fp)

# Cover: editorial_cover header pattern.
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(38)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run("使用指南")
set_run_font(r, size=11, bold=True, color=TEAL)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("专利阅研")
set_run_font(r, size=30, bold=True, color=TEAL_DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
r = p.add_run("从本地文档阅读、图文关联到大模型技术审查")
set_run_font(r, size=15, bold=True, color=INK)

add_figure(
    doc,
    "1-首页.png",
    "图 1  首页：打开本地 DOCX 或 PDF 文件",
    "专利阅研首页，中央提供打开 DOCX 或 PDF 文件按钮。",
    width=6.1,
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("适用于专利代理师、企业知识产权工程师及需要快速理解本地专利稿件的读者")
set_run_font(r, size=10, color=MUTED)

doc.add_page_break()

add_section_heading(doc, "先说结论：这套工具解决什么问题", level=1)
add_body(doc, "专利阅研是一款 Windows 本地端专利阅读辅助工具，主战场是本地 DOCX/PDF。它把专利正文、权利要求和附图放在同一阅读界面中，帮助读者快速完成结构定位、标号—特征映射、图文跳转、批注留痕和整体评级；在 2.x 版本中，还可以接入用户自行配置的大模型与检索服务，辅助识别技术缺陷。")
add_callout(doc, "本文演示范围", "LLM 部分只完整演示“技术理解与技术缺陷”。清楚性/支持性/形式缺陷、新颖性/创造性、确权与维权稳定性三个模块的入口和基本操作相同，本文不重复展开。")

add_section_heading(doc, "建议先记住的四条原则", level=2)
add_bullet(doc, "原文件默认不被覆盖。点击“保存修订版”后，工具另存一份带“修订版”字样的 DOCX/PDF。")
add_bullet(doc, "结构区段、标号映射、OCR 结果和 LLM 结论都属于辅助识别，关键结论需要人工确认。")
add_bullet(doc, "本地阅读与本地 OCR 可不联网；启用云 OCR、LLM 或 Web Search 时，会调用用户自行配置的第三方服务。")
add_bullet(doc, "大模型输出不是法律意见、无效结论或侵权结论。技术事实来源、对比文件和修改建议都应由专业人员复核。")

add_section_heading(doc, "一条最省时的使用路径", level=2)
add_number(doc, "打开 DOCX/PDF，核对四个阅读区段。", quick_path_num)
add_number(doc, "确认“标号—特征”映射，应用到全文与附图。", quick_path_num)
add_number(doc, "在双窗口中对照正文与附图阅读，必要时添加批注。", quick_path_num)
add_number(doc, "完成技术理解、沟通和专利质量评级。", quick_path_num)
add_number(doc, "如需更深入分析，再进入 LLM 审查；先运行技术事实检索，再人工确认外部证据和审查卡片。", quick_path_num)
add_number(doc, "点击“保存修订版”，输出批注版文件与评级记录。", quick_path_num)

add_section_heading(doc, "第一部分　打开文档并确认结构", level=1, page_break=True)
add_body(doc, "点击首页的“打开 DOCX 或 PDF”，选择一份完整的专利稿件。工具会尝试识别说明书摘要、权利要求书、说明书和说明书附图四个区段。不同代理机构的模板、页眉和章节顺序并不统一，因此首次打开时务必花几秒核对。")
add_figure(
    doc,
    "2-结构匹配页.png",
    "图 2  结构匹配页：核对专利四个阅读区段",
    "结构匹配页显示说明书摘要、权利要求书、说明书和说明书附图的识别结果，并允许手动指定。",
)
add_callout(doc, "注意", "“说明书摘要”可能不存在，或被并入其他区段。没有该部分时可以保持未选择；其余区段如识别不准，请使用右侧下拉框手动指定后再开始阅读。", fill=WARN_FILL, label_color=WARN_TEXT)

add_section_heading(doc, "第二部分　理解主界面", level=1, page_break=True)
add_body(doc, "进入主界面后，屏幕被分成三块：左侧是结构导航和正文，中间是全部附图，右侧是标号映射、引用基础判断、批注与评级。正文与附图可以分别滚动，适合边读权利要求、说明书，边核对附图。")
add_figure(
    doc,
    "3-主界面.png",
    "图 3  主界面：正文、附图和审查工具同屏呈现",
    "主界面包含左侧专利正文、中部附图和右侧标号映射与引用基础判断。",
)

add_section_heading(doc, "2.1 结构导航与引用基础判断", level=2)
add_bullet(doc, "点击左侧“说明书摘要”“权利要求书”“说明书”“说明书附图”，可跳转到对应内容。左栏可通过收起图标腾出阅读空间。")
add_bullet(doc, "引用基础判断用于检查“所述、该、上述”等引用词能否沿引用链追溯到首次引入。工具会标出疑似缺乏引用基础的位置，用户需要判断被标记文本是否确实属于特征名称。")
add_bullet(doc, "标题中已经引入的主题名称通常可视为已引入；上下位词、带序号名称和短词包含关系仍需人工辨别。")

add_section_heading(doc, "第三部分　确认标号并阅读附图", level=1, page_break=True)
add_body(doc, "右侧第一步会从说明书附图标记和全文中提取“标号—特征”候选。默认选择出现次数较多或置信度较高的名称；如有歧义，点击特征名称展开候选项，也可以删除误识别标号或通过“+ 自定义补充”增加标号。确认后点击“确认并应用到全文”。")
add_figure(
    doc,
    "4-已识别标号.png",
    "图 4  标号应用后：附图上显示可拖动的浮动标签",
    "附图上的标号已映射为带名称的浮动标签，标签可拖动、关闭，附图可打开缩放窗口。",
)
add_bullet(doc, "浮动标签可以拖动，避免遮挡图纸；误标标签可点击右侧“×”单独关闭。")
add_bullet(doc, "双击附图，或在悬停出现放大提示时点击，可在置顶窗口中缩放查看细节。")
add_bullet(doc, "本地 OCR 使用内置识别能力并结合正文映射表纠错；如字体偏门、标号过小或线条干扰严重，可在顶部 OCR 设置中切换到已配置的云 OCR。")
add_callout(doc, "识别策略", "少量多识别通常比漏识别更容易纠正。建议先保留疑似标号，确认映射后再关闭错误标签；真正缺失的标号可手工补充，再触发 OCR 尝试定位。")

add_section_heading(doc, "第四部分　添加批注并保存修订版", level=1, page_break=True)
add_body(doc, "切换到“审阅”模式，在正文中选中需要批注的精确文字。选区会保持高亮，随后在右侧选择问题类型和程度，填写批注人及批注内容，点击“添加批注”。")
add_figure(
    doc,
    "5-批注.png",
    "图 5  批注流程：选中文字—填写批注—保存修订版",
    "正文中选中文本后，右侧批注区可填写类型、程度、批注人和批注内容。",
)
add_callout(doc, "千万注意", "关闭当前窗口或重新选择文件前，请先点击右上角“保存修订版”。界面中的批注状态不等于已经写回文件；保存后会生成独立修订版，不直接修改原文件。", fill=WARN_FILL, label_color=RED)
add_bullet(doc, "尽量不要跨段选择批注锚点。跨段、重复文本或过长选区会增加定位歧义。")
add_bullet(doc, "若必须批注较长内容，建议拆成若干个语义完整的短选区，便于 WPS 精确定位和后续复核。")
add_bullet(doc, "当前批注写回以 WPS 场景为主；Microsoft Office 等其他软件请以实际打开效果为准。")

add_section_heading(doc, "第五部分　完成整体评级", level=1, page_break=True)
add_body(doc, "右侧底部提供技术理解、沟通和专利质量三个 A—D 评级。评级适合企业内部形成统一的稿件交接记录，也可以帮助后来接手的人员快速判断阅读重点。")
add_figure(
    doc,
    "6-评级区.png",
    "图 6  专利整体评级区",
    "右侧评级区包含技术理解评级、沟通评级和专利质量评级。",
)
add_callout(doc, "留存方式", "点击“保存修订版”时，工具会将评级另存为 Excel 记录。建议企业先统一 A—D 的内部定义，避免不同人员打分尺度不一致。")

add_section_heading(doc, "第六部分　进入 LLM 专利辅助审查", level=1, page_break=True)
add_body(doc, "2.x 版本提供 LLM 辅助审查。可从顶部“LLM 审查”按钮进入，也可从右侧“LLM 专利辅助审查”卡片进入。由于会把选定范围的专利文本发送给用户配置的模型服务商，涉及未公开稿件时请先确认公司保密制度和服务商的数据政策。")
add_figure(
    doc,
    "7-大模型审查入口.png",
    "图 7  LLM 辅助审查的两个入口",
    "主界面顶部和右侧卡片均可进入 LLM 专利辅助审查。",
)

add_section_heading(doc, "6.1 配置模型服务", level=2)
add_body(doc, "在模型设置区选择服务商后，工具会带出默认服务器地址。点击 API Key 旁的“获取”可前往服务商页面申请密钥；填入密钥后点击“连接获取”，读取当前可用模型，并在下拉框中选择。建议选择支持深度思考、长上下文和结构化输出的模型。")
add_figure(
    doc,
    "8-模型设置.png",
    "图 8  模型服务商、服务器地址、API Key 与模型名称",
    "LLM 设置区域允许选择服务商、填写 API Key、获取可用模型并选择模型。",
)
add_callout(doc, "安全提醒", "API Key 只应保存在自己的电脑和自己的账号中，不要截图分享，不要写入示例文件，也不要上传到 GitHub。更换电脑或交付安装包前，应再次检查本地配置。", fill=WARN_FILL, label_color=RED)

add_section_heading(doc, "6.2 四个审查模块如何选择", level=2)
add_body(doc, "LLM 审查提供四个模块。可单选或组合；模型会按选中的模块分批执行，减少长上下文溢出。本文只演示第一个模块。")

table = doc.add_table(rows=5, cols=3)
set_table_geometry(table, [2250, 3150, 3960], indent_dxa=120)
headers = ["模块", "主要输入", "典型输出"]
for i, text in enumerate(headers):
    cell = table.cell(0, i)
    shade_cell(cell, TEAL_LIGHT)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=10, bold=True, color=TEAL_DARK)
rows = [
    ("技术理解与技术缺陷", "专利全文；可选技术事实检索", "作用原理、成立条件、参数边界、工程风险与技术矛盾"),
    ("清楚性、支持性及形式缺陷", "专利全文与内置规则库", "术语、引用、支持、公开充分、形式问题"),
    ("新颖性与创造性", "用户对比文件和/或专利检索结果", "最接近现有技术、区别特征、三步法分析方向"),
    ("确权与维权稳定性", "权利要求、说明书与内置规则库", "稳定性、解释风险、保护范围和取证难点"),
]
for row_idx, row in enumerate(rows, start=1):
    for col_idx, text in enumerate(row):
        cell = table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=INK, bold=(col_idx == 0))

add_figure(
    doc,
    "9-大模型模块介绍.png",
    "图 9  选择技术领域、审查模块、范围和对比文件",
    "LLM 审查设置区展示技术领域、四个审查模块、本次发送范围和对比文件。",
)
add_bullet(doc, "技术领域可以自定义。工具会结合说明书技术领域段落进行识别；如果二者差异明显，应由用户确认。")
add_bullet(doc, "本次发送范围可按全文或区段选择。未公开稿件应遵循最小必要原则，只发送完成本次任务所需内容。")
add_bullet(doc, "新颖性与创造性必须有对比基础：上传对比文件，或开启专利联网检索。没有可靠对比文件时，不应把模型猜测当作新颖性、创造性结论。")

add_section_heading(doc, "第七部分　演示：技术理解与技术缺陷", level=1, page_break=True)
add_body(doc, "这一模块的目标不是找相似专利，而是检索基本理论、工程做法、必要参数、成立条件和失效风险，用外部事实核验主文本中的技术描述是否成立。")

add_section_heading(doc, "7.1 选择技术领域与 Web Search", level=2)
add_body(doc, "勾选“技术理解与技术缺陷”，并在“联网检索”中选择“检索技术事实”。然后配置 Web Search 服务、服务器地址、API Key、搜索引擎和每条检索式返回数量。截图示例使用智谱 Web Search；服务和计费规则以用户自己的账号为准。")
add_figure(
    doc,
    "10-技术判断模块.png",
    "图 10  技术事实检索：技术领域、检索服务和参数",
    "技术判断模块可设置技术领域，勾选检索技术事实，并配置 Web Search 服务。",
)
add_callout(doc, "检索边界", "技术事实检索用于核验原理、参数和工程风险，不作为新颖性或创造性的对比文件检索。专利检索应使用对比文件模块或专利检索 MCP。")

add_section_heading(doc, "7.2 生成并确认技术事实检索方案", level=2, page_break=True)
add_body(doc, "“技术事实检索方案”可以由用户手工填写，也可以点击“由 LLM 生成”。建议先让模型从主文本中提取高风险命题，再按基本理论、工程做法、参数边界和失效风险拆成若干检索问题。")
add_figure(
    doc,
    "11-技术检索llm.png",
    "图 11  点击“由 LLM 生成”形成检索方案",
    "技术事实检索方案文本框和由 LLM 生成按钮。",
)
add_body(doc, "模型生成的检索计划不是最终结论。应删掉与发明点无关、重复、过宽或明显不存在技术争议的问题，保留真正可能影响方案成立的命题。检索方案越聚焦，费用越可控，证据也越容易人工复核。")
add_figure(
    doc,
    "12-大模型判断潜在技术问题，并提供检索计划.png",
    "图 12  人工编辑并确认检索计划",
    "LLM 生成潜在技术问题和检索计划后，用户可人工修改并确认。",
)
add_bullet(doc, "删除明显与发明点无关、重复、过宽或只能得到营销材料的检索命题。")
add_bullet(doc, "对关键参数补充单位、工况、材料、结构对象和比较基准，让检索式具备工程上下文。")
add_bullet(doc, "确认检索计划后再启动联网检索；如果不希望产生费用，可取消联网并让模型仅作文本内部分析。")

add_section_heading(doc, "7.3 运行检索并选择外部证据", level=2, page_break=True)
add_body(doc, "点击“开始辅助审查”后，工具会提示联网检索可能产生额外费用。确认继续后，系统按检索计划分轮调用搜索工具，形成候选外部证据。检索过程中会显示当前阶段、已完成批次和累计卡片数量。")
add_figure(
    doc,
    "13-技术检索内容.png",
    "图 13  检索与审查卡片生成进度",
    "技术事实检索完成后进入人工证据确认，并显示审查卡片生成进度。",
)
add_bullet(doc, "只保留可打开、来源可靠、与待核验命题直接相关的证据。标准、教材、同行评议论文、权威机构资料和厂商技术手册通常比聚合网页更可靠。")
add_bullet(doc, "搜索摘要只能用于初筛。作出专业结论前，应打开来源并阅读必要上下文，核对适用条件、发布日期、试验对象和单位。")
add_bullet(doc, "如果来源打不开、正文缺失或只提供二次转述，应取消勾选，避免把低质量摘要送入后续审查。")

add_section_heading(doc, "7.4 复核技术评判结果", level=2, page_break=True)
add_body(doc, "模型会把主文本中的技术命题、外部事实、缺失条件、潜在风险和修改建议整理为审查卡片。卡片可以按主文本位置、重要程度或问题分类排序；左上角勾选表示保留，未勾选卡片不会写入后续审查报告。")
add_figure(
    doc,
    "14-技术评判结果.png",
    "图 14  技术评判结果：排序、筛选和人工采纳",
    "技术审查卡片展示问题、证据和建议，可排序并由用户决定是否采纳。",
)
add_callout(doc, "判断标准", "优先保留“原文可定位、外部事实可验证、缺失条件明确、建议可操作”的卡片。只表达一般常识、没有来源、超出文本或把可能性写成确定结论的卡片，应退回核验或不采纳。")

add_section_heading(doc, "第八部分　另外三个 LLM 模块如何举一反三", level=1, page_break=True)
add_section_heading(doc, "8.1 清楚性、支持性及形式缺陷", level=2)
add_body(doc, "适合在没有对比文件时先做文本质量体检。重点复核术语一致、引用关系、上位概括、权利要求支持、必要技术特征、公开充分和附图标记。工具内置规则库可以提高问题覆盖率，但最终判断仍应结合具体文本和审查实践。")

add_section_heading(doc, "8.2 新颖性与创造性", level=2)
add_body(doc, "必须以用户上传的对比文件和/或专利联网检索结果为基础。系统会先汇报候选文件及特征覆盖率，由用户选择最接近现有技术，再进行单独对比和创造性三步法分析。不要仅凭普通网页搜索摘要形成法律结论。")

add_section_heading(doc, "8.3 确权与维权稳定性", level=2)
add_body(doc, "从授权后的解释、无效稳定性和取证可行性角度审查权利要求。重点关注功能性限定、使用环境特征、必要技术特征、等同空间、禁止反悔、说明书捐献和可检测性。输出用于提示风险，不替代案件法律意见。")

add_section_heading(doc, "第九部分　推荐的日常工作流", level=1)
add_number(doc, "先用本地阅读功能理解方案：结构定位—标号映射—正文/附图联读。", workflow_num)
add_number(doc, "人工做第一轮批注，记录真正影响理解、支持、清楚性和可实施性的疑点。", workflow_num)
add_number(doc, "如技术原理或参数可疑，启用“技术理解与技术缺陷”，用权威外部资料核验。", workflow_num)
add_number(doc, "如要评价文本质量，启用清楚性/支持性/形式缺陷模块。", workflow_num)
add_number(doc, "只有拿到对比文件或启用专利检索后，再评价新颖性与创造性。", workflow_num)
add_number(doc, "最后从确权与维权角度检查保护范围、稳定性和取证难度。", workflow_num)
add_number(doc, "人工确认卡片和批注，完成评级，然后保存修订版及 Excel 记录。", workflow_num)

add_section_heading(doc, "常见问题与排查", level=1, page_break=True)
add_section_heading(doc, "附图或标号未显示", level=2)
add_bullet(doc, "先核对“说明书附图”区段是否识别正确，再检查图片是否真正嵌入 DOCX。")
add_bullet(doc, "确认“标号—特征”后点击“确认并应用到全文”；必要时手工补充标号。")
add_bullet(doc, "本地 OCR 漏识别时，可提高图像清晰度或切换已配置的云 OCR。")

add_section_heading(doc, "批注没有出现在文件中", level=2)
add_bullet(doc, "确认已经点击“添加批注”，并在退出前点击“保存修订版”。")
add_bullet(doc, "缩短批注锚点，避免跨段选择或选择全文中重复出现的长句。")
add_bullet(doc, "优先使用 WPS 验证批注；其他 Office 软件的兼容性需单独确认。")

add_section_heading(doc, "LLM 没有结果或中途失败", level=2)
add_bullet(doc, "检查模型、API Key、余额、网络和服务商状态。")
add_bullet(doc, "选择支持长上下文和深度思考的模型，并减少一次发送的区段或模块数量。")
add_bullet(doc, "分模块、分数据包执行；失败后使用保留的检索结果和已完成卡片继续，不必重复付费检索。")
add_bullet(doc, "如界面提供诊断信息，可复制不含 API Key 和专利正文的诊断摘要用于排查。")

add_section_heading(doc, "隐私与使用边界", level=1)
add_callout(doc, "本地优先", "DOCX/PDF 解析、正文阅读、批注和本地 OCR 适合在本机完成。只有用户主动配置并启用云 OCR、LLM、Web Search 或 MCP 时，才会调用相应外部服务。")
add_bullet(doc, "未公开专利稿件属于高敏感资料。启用任何云服务前，请确认企业制度、服务条款、数据保留策略和跨境要求。")
add_bullet(doc, "仅上传完成任务所必需的文本或附图；不要在截图、日志、教程或 GitHub 中暴露 API Key。")
add_bullet(doc, "工具输出用于辅助阅读、检索和审查，不构成法律意见、侵权判断、无效结论或授权保证。")

add_section_heading(doc, "获取工具与反馈", level=1)
add_body(doc, "项目仓库：")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
add_hyperlink(p, "https://github.com/talenlin/Patent-Review-Tool", "https://github.com/talenlin/Patent-Review-Tool")
add_body(doc, "如果你在使用中遇到区段识别、标号 OCR、批注写回、模型兼容或检索证据问题，建议记录：工具版本、文件格式、复现步骤、界面截图和不含敏感信息的诊断摘要。这样更容易定位问题，也能推动后续版本改进。")

add_callout(doc, "最后一句", "这款工具的价值不在于替代专业人员，而在于把“结构定位—图文联读—问题留痕—证据核验”串成一条更短、更清楚的阅读路径。先让机器做重复工作，再由专业人员作最终判断。")

doc.core_properties.title = "专利阅研使用指南：从本地文档阅读、图文关联到大模型技术审查"
doc.core_properties.subject = "微信公众号图文使用指导"
doc.core_properties.author = "专利阅研"
doc.core_properties.keywords = "专利阅读, DOCX, PDF, OCR, 批注, LLM, 技术事实检索"

doc.save(OUT_PATH)
print(OUT_PATH)
